#-*- coding: utf-8 -*-
#**********************************************************************
#                 Convert ODK File to Paper Survey
#   Name: Zachary Groff
#   Email: zgroff@poverty-action.org
#   Date: June 17, 2016
#   Purpose: Take ODK files in .xlsx format and convert them to word.
#   Needed Programs: openpyxl, python-docx (can use pip install to install)
#**********************************************************************

#NOTE: You MUST have openpyxl and python-docx installed in order to run this
#program. If you do not have them installed, use pip install to install them.

#***************************************************************
#Specify the following variables. Do not type below the warning.
#***************************************************************

#Specify name of document to import in quotes.
excelname='SPREADSHEET_NAME.xlsx'

#Specify name to save document as in quotes.
wordname='EP_Baseline_Adult_Survey_Master_Modified.docx'

language=''
#Specify language (leave empty if only one choice).
#MUST LEAVE EMPTY IF YOU HAVE ONLY ONE LABEL COLUMN
#MUST SELECT IF YOU HAVE MULTIPLE LANGUAGES IN SURVEY

#Specify the default number of repeat groups.
defaultrc=10

#Specify title of survey if different from "form_title."
formtitle=''

#Stop repeats from showing (=1 to suppress repeats).
suppress=0

#Use tables for formating innermost repeat groups.
tablesinclude=1

#Show relevance.
relevances=1

#Show constraints.
constraints=1

#Show calculate fields (=1 to show calculate fields).
calculates=0

#***********************
#DO NOT TYPE BELOW HERE.
#***********************

#***Import Important Data from ODK File***

#Import docx program.
import openpyxl
import docx
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE
from docx.shared import Pt
from docx.shared import Inches
import string
import unicodedata

#Import excel document.
wb = openpyxl.load_workbook(excelname)
survey=wb['survey']
choices=wb['choices']
settings=wb['settings']

#Create word document for output of program.
document = Document()

#Create a dictionary for each sheet linking column titles to the letter of the column.
#For example, if column A is type, survcoldict[type]==A.
survcoldict={}
chcoldict={}
for l in string.ascii_uppercase:
    #Create a dictionary for survey sheet.
    #Enter column titles that are strings into dictionary.
    if isinstance(survey[l+'1'].value, str):
        survcoldict[str(survey[l+'1'].value)]=l
    #Enter column titles that are unicode into dictionary.
    elif survey[l+'1'].value!=None:
        survcoldict[unicodedata.normalize('NFKD', survey[l+'1'].value).encode('ascii','ignore')]=l
    #Create a dictionary for choices sheet.
    if isinstance(choices[l+'1'].value, str):
        chcoldict[str(choices[l+'1'].value)]=l        
    elif choices[l+'1'].value!=None:
        chcoldict[unicodedata.normalize('NFKD', choices[l+'1'].value).encode('ascii','ignore')]=l
    #Add entries to dictionary for "label" and "hint" for specified language for surveys with multiple languages.
    #This makes it so that "label" has an entry even if the column choices are "label:English" and "label:Spanish."
    if survey[l+'1'].value=='label:'+language:
        survcoldict['label']=l
    if choices[l+'1'].value=='label:'+language:
        chcoldict['label']=l
    if survey[l+'1'].value=='hint:'+language:
        survcoldict['hint']=l
    if choices[l+'1'].value=='hint:'+language:
        chcoldict['hint']=l

#Raise an error if there is no language entered but there should be.
if "label" not in survcoldict.keys():
    print("A language is not specified, but the survey has either multiple languages or has not found a 'label' column.")
    raise NameError("No Label")


print "The columns in the survey tab are: "+str(survcoldict)
print "The columns in the choices tab are: "+str(chcoldict)

#Add title and initial info to word document.
print "Title of form: "+str(settings['A2'].value)
if formtitle=='':
    formtitle=settings['A2'].value
document.add_heading(formtitle, 0)
intro1=document.add_paragraph('Start time:__________                                End time:__________')
intro2=document.add_paragraph('Device ID:____________________               Subscriber ID:____________________')
intro3=document.add_paragraph('Sim ID (Serial):____________________    Device phone number:____________________')


#***Define Functions for Program***

#Define printing function for questions not in a repeat group.
def QuestionState(query, tip, sort, chart, colnum, tableyesno):
    #
    #----------------------------------------------------------------------------------------------
    #query=Text of the question to be printed.
    #tip=Hint for the question (may be None if blank).
    #sort=Type of question (e.g. select one, text).
    #chart=The current table, if the document is currently printing a table.
    #colnum=The column number for the current table if the document is currently printing a table.
    #tableyesno=The answer to the question "Is the document currently printing a table?"
    #----------------------------------------------------------------------------------------------
    #
    #Define question number global to update over the course of the program.
    global number
    number += 1
    #Add question (query) to document, with hint (tip) if relevant.
    #The question format varies depending on whether it is in a table or not. Start by defining the way to print questions not in tables.
    if tableyesno==0:
        qp=document.add_paragraph(str(number)+'. '+query)
        qp.paragraph_format.space_after = Pt(0)
        qp.paragraph_format.space_before = Pt(12)
        if tip is not None:
            qph=document.add_paragraph('')
            qph.add_run(tip).italic = True
    #Now, add the question to the document if the table format is needed.
    else:
        #1. Write out the text of the question
        if tip!=None:
            fullquestion=query+' ('+tip+')'
        else:
            fullquestion=query
        #2. If the question is a select_one or select_multiple question, include that instruction and the title of the choices list in the question.
        if sort.partition(' ')[0]=='select_one':
            chart.cell(0, colnum).text=str(number)+'. SELECT ONE ('+sort.partition(' ')[2]+'): '+fullquestion
        elif sort.partition(' ')[0]=='select_multiple':
            chart.cell(0, colnum).text=str(number)+'. SELECT MULTIPLE ('+sort.partition(' ')[2]+'): '+fullquestion
        #3. If the question is not select_one or select_multiple, print it without that.
        else:
            chart.cell(0, colnum).text=str(number)+'. '+fullquestion

#Define function to print the choice options for a question.
def OptionList(paths, category):
    #
    #----------------------------------------------------------------------------------------------
    #paths=Title of the choices list, i.e. list_name in the ODK file.
    #category=select_one or select_multiple.
    #----------------------------------------------------------------------------------------------
    #
    #Go row by row through the choices spreadsheet.
    for y in range(2, choices.max_row+1):
        #Store list_name, value, and label columns as variables ("value" as "name" to avoid confusion).
        list_name=choices[chcoldict['list_name']+str(y)].value
        name=choices[chcoldict['value']+str(y)].value
        #Fix label variable to correctly reference previous variables and have the right format.
        label=ReplaceRefs(choices[chcoldict['label']+str(y)].value, 'C')
        if not isinstance(label, (str, unicode)):
            label=unicode(str(label), 'utf-8')
        #Print name and label in row one way if it's select_multiple and another way if it's select_one.
        if list_name==paths:
            if category=='select_multiple':
                op=document.add_paragraph('__ '+str(name)+' - '+unicodedata.normalize('NFKD', label).encode('ascii', 'ignore'))
            else:
                op=document.add_paragraph('[] '+str(name)+' - '+unicodedata.normalize('NFKD', label).encode('ascii', 'ignore'))
            op.paragraph_format.space_after=Pt(0)
            op.paragraph_format.left_indent=Inches(0.5)
        
#Define a function that returns 1 if a begin repeat should begin a table and 0 otherwise.
#A begin repeat should begin a table if and only if it is the innermost repeat group.
#This program looks row by row through the survey, bookmarking where the current repeat group starts.
#When the program finds where the repeat group ends, if there has not been another repat group in betwen, it starts a table.
def TableTime(repeatgroup):
    #
    #----------------------------------------------------------------------------------------------
    #repeatgroup=The name of a repeat group (the one we are currently on).
    #----------------------------------------------------------------------------------------------
    #
    #The default is that it is not a table.
    tabletime=0
    #Go row by row through the survey spreadsheet.
    for x in range(8, survey.max_row+1):
        #Set tabletime to 1 once we reach the row where the repeat group starts.
        if survey[survcoldict['type']+str(x)].value=='begin repeat' and survey[survcoldict['name']+str(x)].value==repeatgroup:
            tabletime=1
        #If another repeat group starts after the beginning of this repeat group but before we reach the end of this repeat group, return 0—this is not the innermost repeat group.
        if survey[survcoldict['type']+str(x)].value=='begin repeat' and survey[survcoldict['name']+str(x)].value!=repeatgroup:
            if tabletime==1:
                return 0
        #If we reach the end of this repeat group before another repeat group starts, return the current value of tabletime (1).
        if survey[survcoldict['type']+str(x)].value=='end repeat':
            return tabletime

#Generate function that scans a phrase and replaces references to earlier fields.
def ReplaceRefs(phrase, mode):
    #
    #--------------------------------------------------------------------------------------------------------
    #phrase=The text of the field where we are trying to replace references.
    #mode=The type of text: 'Q' for question, 'A' for relevance/calculation/constraint, and 'C' for choice.
    #--------------------------------------------------------------------------------------------------------
    #
    tempphrase=phrase
    if isinstance(tempphrase, str) or isinstance(tempphrase, unicode):
        #Identify phrases that make references to earlier fields.
        if '${' in tempphrase and '}' in tempphrase:
            #Initiate reference replacements.
            #referring is a constant that indicates when we are in the process of scanning the string for the name of the previous field.
            referring=0
            #ref is the name of the previous field being referred to.
            ref=''
            #replacements is the dictionary of previous fields referred to and the text to replace them.
            replacements={}
            #Go letter by letter through the phrase.
            for n in range(1, len(tempphrase)):
                if tempphrase[n]=='}':
                    #If we have reached the end of the name of the field, change referring to 0.
                    referring=0
                    #If we have reached the end of the name of the field and the field referred to has a number, add a text reference to the dictionary of replacements.
                    if ref in qnumbers and mode=='Q':
                        replacements['${'+ref+'}']=' _________ (Answer to Q'+str(qnumbers[ref])+') '
                    if ref in qnumbers and mode=='A':
                        replacements['${'+ref+'}']= ' the answer to Q'+str(qnumbers[ref])+' '
                    if ref in qnumbers and mode=='C':
                        replacements['${'+ref+'}']= '[Answer to Q'+str(qnumbers[ref])+']'
                    #Reset ref to blank once we have reached the end of the field and made a replacement.
                    ref=''
                #If we are in the process of scanning for the name of a previous field and have not reached the end of the field name, add the current letter onto the string.
                if referring==1:
                    ref=ref+tempphrase[n]
                #Start scanning at the letters '${' for the name of a previous field.
                if tempphrase[n-1]=='$' and tempphrase[n]=='{':
                    referring=1
            #Go through the dictionary and make all the proper replacements.
            for key in replacements.keys():
                tempphrase=tempphrase.replace(key, replacements[key])
    return tempphrase     

#Generate function to translate ODK expressions into English.
def TranslateCalc(exp, variety):
    #
    #--------------------------------------------------------------------------------------------------------
    #exp=The text of the expression.
    #variety=The type of field.
    #--------------------------------------------------------------------------------------------------------
    #
    newexp=exp
    #Set "verb" equal to "is" or "must be" depending on whether it is a constraint or not.
    if variety!='constraint':
        verb='is '
    else:
        verb='must be '
    #Replace functions with English phrases that have the same meaning.
    newexp=newexp.replace('selected(', 'selected options include [')
    newexp=newexp.replace('string-length(', 'length of ')
    newexp=newexp.replace('.', 'answer ')
    newexp=newexp.replace('>=', verb+'greater than or equal to ')
    newexp=newexp.replace('<=', verb+'less than or equal to ')
    newexp=newexp.replace('>', verb+'greater than ')
    newexp=newexp.replace('<', verb+'less than ')
    newexp=newexp.replace('+', 'plus ')
    newexp=newexp.replace('-', 'minus ')
    newexp=newexp.replace('/', 'divided by')
    newexp=newexp.replace('*', 'times ')
    newexp=newexp.replace('!=', 'does not equal ')
    newexp=newexp.replace('=', 'equals ')
    #Replace paranetheses with brackets.
    if '(' in newexp:
        newexp=newexp.replace(')', '] ')
    newexp=newexp.replace('(', ' [').strip()
    #Capitalize the first letter if the expression is not a relevance function (in which case it starts with the name of a field, which should not be capitalized). 
    if variety!='relevance':
        newexp=newexp[0].upper() + newexp[1:]
    #Add a period to the end for proper punctuation.
    newexp=newexp+'.'
    return newexp

#Now turn the survey into a word document, row by row.
#repeat tells us how many repeat groups a row is nested inside.
repeat=0
#number is the number of the question.
number=0
#qnumbers is a dictionary whose keys are the names of fields in the ODK file and whose values are the corresponding numbers.
#Fields in repeat groups correspond to multiple numbers unless the repeat group is a table, so the value is a list of numbers separated by "/".
qnumbers={}
def Program(a, b, roundnum, tableyesno=0, repeat=0, repeatcount=0):
    #
    #--------------------------------------------------------------------------------------------------------
    #a=The number of the first row to be used in the ODK file.
    #b=The number of the last row to be used in the ODK file.
    #roundnum=The current iteration of a repeat group as a string.
    #tableyesno=1 if the program should output a table, 0 if not.
    #repeat=How many repeat groups the execution of the Program() function is nested inside.
    #repeatcount=The number of times to repeat a repeat group.
    #--------------------------------------------------------------------------------------------------------
    #
    #skip_until marks the end of a repeat group so that a repeat group is not processed an extra time after completion.
    #Initiate skip_until as 0 and update skip_until when processing a repeat group.
    skip_until=0
    #Go row by row through the survey.
    for x in range(a, b+1):
        #If skip_until has been set, do not print any rows before the end of the repeat group (skip_until).
        if x<=skip_until:
            continue
        #Set variable type equal to the type of question in row x.
        typ=''
        if survey[survcoldict['type']+str(x)].value!=None:
            typ=unicodedata.normalize('NFKD', survey[survcoldict['type']+str(x)].value).encode('ascii', 'ignore')
        if typ=='':
            print "\nSurvey appears to skip a row at line "+str(x)+" because question type is blank. Please make sure this is correct."
        #Create a variable, programmed, that indicates that the current row is of a type that this program can process.
        programmed=typ in ['text', 'integer', 'geopoint', 'note', 'begin group', 'end group', 'begin repeat', 'end repeat'] or typ.partition(' ')[0] in ['select_one', 'select_multiple']
        #Add the current question number to the dictionary qnumbers.
        if survey[survcoldict['name']+str(x)].value not in qnumbers.keys():
            qnumbers[survey[survcoldict['name']+str(x)].value]=number+1
        #If there is already a question number for this row (in a repeat group, for example), add the number to the value as part of a list separated by "/".
        else:
            qnumbers[survey[survcoldict['name']+str(x)].value]=str(qnumbers[survey[survcoldict['name']+str(x)].value])+"/"+str(number+1)
        #Set question equal to the value of the "label" column (for the specified language if relevant).
        question=survey[survcoldict['label']+str(x)].value
        #If the row is a calculate field, set question equal to the "calculation" column.
        if typ=='calculate' or typ=='calculate_here':
            question=survey[survcoldict['calculation']+str(x)].value
        #Remove unprintable characters from the question.
        if question!=None:
            question=unicodedata.normalize('NFKD', question).encode('ascii', 'ignore')
        #Replace references to other fields.
        if question!=None and programmed:
            question=ReplaceRefs(question, 'Q')
        if (typ=='calculate' or typ=='calculate_here') and question!=None:
            question=TranslateCalc(ReplaceRefs(question, 'A'), 'calculation')
        #Make question an empty string rather than NoneType if relevant.
        if question==None:
            question=''

        #Indicate current group and question in Python Shell.
        if programmed and typ not in ['begin group', 'end group', 'begin repeat', 'end repeat']:
            print "\nProcessing question number "+str(number+1)+":"
            print "Question—"+question
        if typ=='begin group':
            print "\n*****************New Group: "+question+"*****************"
        if typ=='end group':
            print "\n*****************End Group: "+question+"*****************"
    
        #Create a variable, hint, equal to the value of the "hint" column.
        hint=survey[survcoldict['hint']+str(x)].value
        if isinstance(hint, unicode):
            hint=unicodedata.normalize('NFKD', hint).encode('ascii', 'ignore')
            hint=str(hint)
            hint=hint+'.'

        #Create a variable, constraint, equal to the value of the "constraint" column.
        constraint=survey[survcoldict['constraint']+str(x)].value
        if isinstance(constraint, unicode):
            unicodedata.normalize('NFKD', constraint).encode('ascii', 'ignore')
            constraint=str(constraint)
        #Replace references and mathematical symbols in constraint, and add constraint to the hint if constraints is specified above.
        if isinstance(constraint, str):
            constraint=TranslateCalc(ReplaceRefs(constraint, 'A'), 'constraint')
            if isinstance(hint, str) and constraints==1:
                hint=hint+'  '+constraint
            elif constraints==1:
                hint=constraint
                
        #Create a variable, relevance, equal to the value of the "relevance" column.                
        relevance=survey[survcoldict['relevance']+str(x)].value
        if isinstance(relevance, unicode):
            unicodedata.normalize('NFKD', relevance).encode('ascii', 'ignore')
            relevance=str(relevance)
        #Replace references and mathematical symbols in relevance and add relevance to the hint if relevances is specified above.
        if isinstance(relevance, str):
            relevance=TranslateCalc(ReplaceRefs(relevance, 'A'), 'relevance')
            if isinstance(hint, str) and relevances==1:
                hint=hint+'  Only ask if '+relevance
            elif relevances==1:
                hint='Only ask if '+relevance
        if isinstance(hint, str):
            hint=hint.replace('..', '.')

        #The way repeat groups operate in this program is that for each repeat group, the program first prints a heading (i.e. "Family Survey").
        #Then, the program processes each question in the repeat group repeatedly (e.g. with the heading "Family Survey: Round 1").
        #Each time the program processes a begin repeat, it adds one to the value of repeat (and subtracts one when the repeat group ends).
        #Because the program processes each begin repeat group twice (once for the overall heading and once for the actual repeat), repeat is always an even number when the program starts an entirely new repeat group.
        #This means that repeat % 2==0 if and only if the program is processing a repeat group for the first time.
        if typ=='begin repeat':
            #Indicate current repeat group in Python Shell.
            print "\n*************New Repeat Group: "+question+"**************"
            #Set repeatcount based on the repeat_count column for row x.
            repeatcount=survey[survcoldict['repeat_count']+str(x)].value
            #If no repeat count is specified, use the default number of repeats from above (defaultrc).
            if not isinstance(survey[survcoldict['repeat_count']+str(x)].value, (int, long)):
                repeatcount=defaultrc
            #If this is the first time processing the repeat group, print an overall heading (e.g. Family Survey; see note above).
            if repeat % 2==0 and suppress==0:
                rtitle=document.add_heading('', 2)
                rtitle.add_run(question).underline=True
            #Otherwise, print a subheading for this round of the repeat (e.g. Family Survey: Round 1).
            else:
                document.add_heading(question+roundnum, 2)
            #Update repeat to mark that we are inside a repeat group.
            repeat=repeat+1
            #Create a temporary repeat variable (check) to figure out how long the repeat group is.
            check=repeat
            #We want to set d equal to the last row of the repeat group.
            d=0
            #Go through the remainder of the survey.
            for z in range(x+1, b+1):
                #If we reach an end repeat and have gone through an even number of begin and end repeats in between, then we know this is the end of the current repeat group.
                #Therefore, in this case, set d equal to the row of this end repeat.
                if survey[survcoldict['type']+str(z)].value=='end repeat' and check==repeat:
                    d=z
                    break
                #For each new begin repeat, add one to check to note that one of the next end repeats is an end to this inner repeat group.
                elif survey[survcoldict['type']+str(z)].value=='begin repeat':
                    check=check+1
                #For each end repeat that is not the final one, subtract one to note that we are out of an inner repeat group.
                elif survey[survcoldict['type']+str(z)].value=='end repeat':
                    check=check-1
            #If we are including tables, determine if this repeat group should be the start of a table.
            if tablesinclude==1 and suppress==0:
                tableyesno=TableTime(unicodedata.normalize('NFKD', survey[survcoldict['name']+str(x)].value).encode('ascii', 'ignore'))
            #If this repeat group should be the start of the table, do the following.
            if tableyesno==1:
                #Set the number of rows for the table equal to either the current repeat count or the default (defaultrc).
                if repeatcount!=None:
                    rowcount=repeatcount+1
                else:
                    rowcount=defaultrc
                #Create a grid-shaped table with this number of rows and no columns.
                table=document.add_table(rows=rowcount, cols=0)
                table.style='Table Grid'
                #Make the table able to take any size.
                table.autofit=False
                #Make the first column a number column with a number for each iteration. 180000 seems like the appropriate size for this.
                table.add_column(180000)
                table.cell(0, 0).text='#'
                for n in range(1, rowcount):
                    table.cell(n, 0).text=str(n)+'.'
                #Create a variable, newcol, that indicates the number of the table column we are currently editing.
                newcol=0
                #Create a dictionary, typedict, that keeps track of the type of question in each column of the table.
                typedict={}

        if typ=='end repeat':
            #Indicate end of current repeat group in Python Shell.
            print "\n*************End Repeat Group: "+question+"**************"
            #Update repeat to indicate the current repeat group has ended.
            repeat=repeat-1
            #If this repeat group is a table, do the following.
            if tableyesno==1:
                #Now, make the table fit the page automatically.
                table.autofit=True
                #For each multiple-choice question in the table, do the following:
                for opt in typedict:
                    options=opt
                    choicetype=typedict[options]
                    #Add a blank line, underlined.
                    o=document.add_paragraph('')
                    o.add_run(options).underline = True
                    o.paragraph_format.space_before=Pt(12)
                    #List the choice options.
                    OptionList(options, choicetype)
                #Add a blank line.
                document.add_paragraph('')
            #Now note that we are done with the table.
            tableyesno=0

        #If we are currently not printing a table, print the questions this way:
        if tableyesno==0:
            #Add a heading for a begin group.
            if typ=='begin group':
                document.add_heading(question, 1)
            #Print the question and then a long blank line to write the answer if the question is text, calculate, or calculate_here.
            if typ=='text' or ((typ=='calculate' or typ=='calculate_here') and calculates==1):
                QuestionState(question, hint, typ, '', '', tableyesno)
                ans=document.add_paragraph('_________________________________________________________________________________________________________')
                ans.paragraph_format.space_before=Pt(6)
            #Print the question and then a short blank line to write the answer if the question is numeric.
            if typ=='integer' or typ=='decimal':
                QuestionState(question, hint, typ, '', '', tableyesno)
                ans=document.add_paragraph('__________________________')
                ans.paragraph_format.space_before=Pt(6)
            #Print the question, instructions, and then the list of options if the question is a select_one.
            if typ.partition(' ')[0]=='select_one':
                question=question+' (select one)'
                QuestionState(question, hint, typ, '', '', tableyesno)
                choicetype=typ.partition(' ')[0]
                options=typ.partition(' ')[2]
                OptionList(options, choicetype)
            #Print the question, instructions, and then the list of options if the question is a select_multiple.
            if typ.partition(' ')[0]=='select_multiple':
                question=question+' (select multiple)'
                QuestionState(question, hint, typ, '', '', tableyesno)
                choicetype=typ.partition(' ')[0]
                options=typ.partition(' ')[2]
                OptionList(options, choicetype)
            #Print the note if the row is a note.
            if typ=='note':
                QuestionState(question, hint, typ, '', '', tableyesno)
            #Print the question and then blanks for longitude, latitude, altitude and accuracy to match ODK style if the question is a geopoint.
            if typ=='geopoint':
                QuestionState(question, hint, typ, '', '', tableyesno)
                ans=document.add_paragraph("Latitude:  __  __* __' __\"")
                ans=document.add_paragraph("Longitude: __  __* __' __\"")
                ans=document.add_paragraph("Altitude:  ______m")
                ans=document.add_paragraph("Accuracy:  ______m")
        #If we are currently printing a table, print the questions this way:
        else:
            #Add a blank column with the name of the group for a begin group.
            if typ=='begin group':
                table.add_column(914400)
                table.cell(0, newcol).text='BEGIN GROUP: '+question
            #Add a blank column noting the end of a group for an end group.
            if typ=='end group':
                table.add_column(914400)
                table.cell(0, newcol).text='END GROUP'
            #Add a blank column with the question or note on top if the row is text, note, calculate, or calculate here (and the user has asked to print calculates if relevant).
            if typ=='text' or (typ=='note') or ((typ=='calculate' or typ=='calculate_here') and calculates==1):
                table.add_column(914400)
                QuestionState(question, hint, typ, table, newcol, tableyesno)
            #Add a small blank column with the question on top if the row is a numeric question.
            if typ=='integer' or typ=='decimal':
                table.add_column(360000)
                QuestionState(question, hint, typ, table, newcol, tableyesno)
            #Add a blank column with the question on top if the row is a select_one question.
            if typ.partition(' ')[0]=='select_one':
                table.add_column(914400)
                QuestionState(question, hint, typ, table, newcol, tableyesno)
                typedict[typ.partition(' ')[2]]='select_one'
            #Add a blank column with the question on top if the row is a select_multiple question.
            if typ.partition(' ')[0]=='select_multiple':
                table.add_column(914400)
                QuestionState(question, hint, typ, table, newcol, tableyesno)
                typedict[typ.partition(' ')[2]]='select_multiple'
            #Add a large blank column with the question on top and space for lattitude, longitude, altitude, and accuracy if the row is a geopoint question.
            if typ=='geopoint':
                table.add_column(1600000)
                QuestionState(question, hint, typ, table, newcol, tableyesno)
                for n in range(1, rowcount):
                    table.cell(n, newcol).text="Latitude:  __  __* __' __\" \n Longitude: __  __* __' __\" \n Altitude:  ______m \n Accuracy:  ______m"
            #Update newcol to indicate that we are one column over.
            if programmed:
                newcol=newcol+1
                
        #If we have only printed the heading of a repeat group (which is true if and only if repeat % 2==1; see note above re: begin repeat) and are in a begin repeat that does not start a table, start repeating the group, assuming we are printing repeats.
        if repeat % 2==1 and suppress==0 and typ=="begin repeat" and tableyesno==0:
            #Set skip_until to the last row of the repeat group so we do not print any of the repeat group after it is done repeating.
            skip_until=d
            c=x
            #Process the repeat group [repeatcount] number of times.
            for i in range(0, repeatcount):
                Program(c, d, ': Round '+str(i+1), 0, repeat)
            #After repeating the repeat group [repeatcount] number of times, add an extra space.
            document.add_paragraph('')

#Run the Program() function over the entire ODK file.
Program(8, survey.max_row, '', 0, 0)

#Save the output as [wordname].        
document.save(wordname)
